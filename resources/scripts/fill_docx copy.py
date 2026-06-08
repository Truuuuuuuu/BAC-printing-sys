import sys
import json
import copy
import re
from docx import Document
from docx.oxml.ns import qn

SPACE_ATTR = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space'


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 5:
        sys.exit("Usage: fill_docx.py <template> <args_json> <table_rows_json> <output>")

    template_path = sys.argv[1]
    output_path   = sys.argv[4]

    try:
        args = json.loads(sys.argv[2])
        if not isinstance(args, dict):
            args = {}
    except json.JSONDecodeError as e:
        sys.exit(f"Invalid args JSON: {e}")

    try:
        table_rows = json.loads(sys.argv[3])
        if not isinstance(table_rows, dict):
            table_rows = {}
    except json.JSONDecodeError as e:
        sys.exit(f"Invalid table_rows JSON: {e}")

    try:
        doc = Document(template_path)
    except Exception as e:
        sys.exit(f"Could not open template: {e}")

    # Replace in top-level body paragraphs (outside tables)
    for para in doc.paragraphs:
        replace_in_paragraph(para, args)

    # Replace in headers and footers
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            if hf is not None:
                for para in hf.paragraphs:
                    replace_in_paragraph(para, args)

    # Process every table
    for table in doc.tables:
        process_table(table, args, table_rows)

    try:
        doc.save(output_path)
    except Exception as e:
        sys.exit(f"Could not save output: {e}")


# ── Paragraph replacement ─────────────────────────────────────────────────────

def format_value(key: str, val) -> str:
    """Apply suffix-based casing to a substitution value."""
    if val is None:
        return ''
    val = str(val)
    if key.endswith('_upper'):
        return val.upper()
    if key.endswith('_lower'):
        return val.lower()
    if key.endswith('_capitalize'):
        return val.capitalize()
    return val


def replace_in_paragraph(para, context: dict) -> None:
    """
    Replace all {{key}} placeholders in a paragraph.

    Handles two layouts Word produces:
      - Intact:  each placeholder sits entirely within one run → replace in-place
                 (preserves individual run formatting like bold/italic/font)
      - Split:   Word has broken the placeholder across multiple runs → merge all
                 runs into runs[0], clear the rest
                 (loses per-run formatting within the placeholder span, unavoidable)
    """
    if not para.runs:
        return

    full_text = ''.join(r.text for r in para.runs)

    # Quick exit — nothing to do
    if '{{' not in full_text:
        return

    # Determine which placeholders are actually present
    active = {k: v for k, v in context.items() if ('{{' + k + '}}') in full_text}
    if not active:
        return

    # Check whether every active placeholder is intact in exactly one run
    all_intact = all(
        any(('{{' + k + '}}') in r.text for r in para.runs)
        for k in active
    )

    if all_intact:
        for run in para.runs:
            for key, val in active.items():
                ph = '{{' + key + '}}'
                if ph in run.text:
                    run.text = run.text.replace(ph, format_value(key, val))
    else:
        # Merge strategy: rebuild full text, write to first run, blank the rest
        result = full_text
        for key, val in active.items():
            result = result.replace('{{' + key + '}}', format_value(key, val))
        para.runs[0].text = result
        for run in para.runs[1:]:
            run.text = ''


# ── Table processing ──────────────────────────────────────────────────────────

PLACEHOLDER_RE = re.compile(r'\{\{row_([a-z]+)_')


def row_group(row) -> str | None:
    """Return the group letter of the first dynamic placeholder found in a row."""
    for cell in row.cells:
        for para in cell.paragraphs:
            text = ''.join(r.text for r in para.runs)
            m = PLACEHOLDER_RE.search(text)
            if m:
                return m.group(1)
    return None


def process_table(table, args: dict, table_rows: dict) -> None:
    """
    Process one table:
    - Static rows (no {{row_X_}} placeholders) → replace {{key}} in-place
    - Template rows (contain {{row_X_}} placeholders) → clone once per data row,
      fill each clone, remove the original template row

    Supports multiple distinct template-row groups in the same table (e.g. both
    'a' rows and 'b' rows living in one table).
    """
    # Index every row by group (None = static)
    row_meta = [(i, row_group(row)) for i, row in enumerate(table.rows)]

    # Collect the distinct groups present in this table
    groups_seen = dict()  # group -> first template row index (for ordering)
    for i, g in row_meta:
        if g is not None and g not in groups_seen:
            groups_seen[g] = i

    if not groups_seen:
        # Purely static table
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, args)
        return

    # We will rebuild the table's row list in-place via XML manipulation.
    # Build a plan: for each existing row, decide what to emit.
    parent = table._tbl
    existing_trs = list(parent)  # snapshot before we touch anything

    # Map row index → original <w:tr> element
    tr_by_index = {i: tr for i, tr in enumerate(existing_trs) if tr.tag.endswith('}tr')}

    # Build the ordered replacement list
    replacement: list = []  # list of <w:tr> elements to write back

    # Track which template-row indices we've already expanded so we don't
    # emit them again if the same group appears on multiple rows (unusual but safe).
    emitted_groups: set = set()

    for i, g in row_meta:
        tr = tr_by_index.get(i)
        if tr is None:
            continue  # non-<w:tr> element (e.g. table properties) — handled separately

        if g is None:
            # Static row: replace placeholders and keep as-is
            for cell in table.rows[i].cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, args)
            replacement.append(tr)
        else:
            if g in emitted_groups:
                # Second template row for the same group — drop it
                # (the expansion was already inserted by the first occurrence)
                continue
            emitted_groups.add(g)

            data_rows = table_rows.get(g, [{}])  # fall back to one blank row
            for row_data in data_rows:
                new_tr = copy.deepcopy(tr)
                fill_tr(new_tr, {**args, **row_data})
                replacement.append(new_tr)
            # Original template row is intentionally not appended → removed

    # Rewrite the parent's child list: remove all existing <w:tr> elements
    # then insert the replacement list at the position of the first <w:tr>.
    first_tr_pos = next(
        (idx for idx, child in enumerate(list(parent)) if child.tag.endswith('}tr')),
        0
    )
    for tr in list(parent):
        if tr.tag.endswith('}tr'):
            parent.remove(tr)

    for offset, new_tr in enumerate(replacement):
        parent.insert(first_tr_pos + offset, new_tr)


def fill_tr(tr_elem, context: dict) -> None:
    """Fill all {{placeholders}} in a raw <w:tr> XML element."""
    for cell_elem in tr_elem.findall('.//' + qn('w:tc')):
        for para_elem in cell_elem.findall('.//' + qn('w:p')):
            runs = para_elem.findall('.//' + qn('w:r'))
            if not runs:
                continue

            full_text = ''.join(
                (r.find(qn('w:t')).text or '') if r.find(qn('w:t')) is not None else ''
                for r in runs
            )

            if '{{' not in full_text:
                continue

            result = full_text
            for key, val in context.items():
                result = result.replace('{{' + key + '}}', format_value(key, val))

            first_t = runs[0].find(qn('w:t'))
            if first_t is not None:
                first_t.text = result
                first_t.set(SPACE_ATTR, 'preserve')

            for run in runs[1:]:
                t = run.find(qn('w:t'))
                if t is not None:
                    t.text = ''


if __name__ == '__main__':
    main()