import sys
import json
import copy
import re
from docx import Document
from docx.oxml.ns import qn

doc = Document(sys.argv[1])
args = json.loads(sys.argv[2])
table_rows = json.loads(sys.argv[3]) if len(sys.argv) > 4 else {}

if not isinstance(args, dict):
    args = {}
if not isinstance(table_rows, dict):
    table_rows = {}

def format_value(key, val):
    if val is None:
        return ''
    val = str(val)
    if key.endswith('_upper'):
        return val.upper()
    if key.endswith('_lower'):
        return val.lower()
    if key.endswith('_capitalize'):
        return val.capitalize()
    if key.endswith('_date'):     
        try:
            from datetime import datetime
            return datetime.strptime(val, '%Y-%m-%d').strftime('%B %d, %Y')
        except ValueError:
            return val
        
    if key.endswith('_formatAmount'):
        try:
            return '{:,.2f}'.format(float(val.replace(',', '')))
        except (ValueError, TypeError):
            return val
        
    if key.endswith('_wordNum'):                   
        try:
            import num2words
            n = int(val.replace(',', ''))
            words = num2words.num2words(n)
            return '{} ({})'.format(words, n)
        except (ValueError, TypeError, ImportError):
            return val
        
    if key.endswith('_numWord'):                   
        try:
            from num2words import num2words
            n = int(val.replace(',', ''))
            return num2words(n)
        except (ValueError, TypeError, ImportError):
            return val
    return val    
    

def merge_and_replace_paragraph(para, extra=None):
    if not para.runs:
        return
    context = {**args, **(extra or {})}
    full_text = ''.join(run.text for run in para.runs)

    has_placeholder = any(('{{' + key + '}}') in full_text for key in context)
    if not has_placeholder:
        return

    all_intact = all(
        any(('{{' + key + '}}') in run.text for run in para.runs)
        for key in context
        if ('{{' + key + '}}') in full_text
    )

    if all_intact:
        for run in para.runs:
            for key, val in context.items():
                placeholder = '{{' + key + '}}'
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, format_value(key, val))
    else:
        for key, val in context.items():
            placeholder = '{{' + key + '}}'
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, format_value(key, val))
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ''

def get_row_group(row):
    """Detect which row_X_ group prefix this template row uses e.g. 'a', 'b'"""
    for cell in row.cells:
        for para in cell.paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            match = re.search(r'\{\{row_([a-z])_', full_text)
            if match:
                return match.group(1)
    return None

# Static body paragraphs
for para in doc.paragraphs:
    merge_and_replace_paragraph(para)

# Tables
for table in doc.tables:
    template_row_idx = None
    group = None

    for i, row in enumerate(table.rows):
        g = get_row_group(row)
        if g:
            template_row_idx = i
            group = g
            break

    if template_row_idx is None:
        # No dynamic rows — static replacement only
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    merge_and_replace_paragraph(para)
        continue

    rows_data = table_rows.get(group, [{}])
    template_tr = table.rows[template_row_idx]._tr
    new_trs = []

    for row_data in rows_data:
        new_tr = copy.deepcopy(template_tr)
        for cell_elem in new_tr.findall('.//' + qn('w:tc')):
            for para_elem in cell_elem.findall('.//' + qn('w:p')):
                runs = para_elem.findall('.//' + qn('w:r'))
                if not runs:
                    continue
                full_text = ''.join(
                    (r.find(qn('w:t')).text or '')
                    if r.find(qn('w:t')) is not None else ''
                    for r in runs
                )
                context = {**args, **row_data}
                for key, val in context.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in full_text:
                        full_text = full_text.replace(placeholder, format_value(key, val))
                first_t = runs[0].find(qn('w:t'))
                if first_t is not None:
                    first_t.text = full_text
                    first_t.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', 'preserve')
                for run in runs[1:]:
                    t = run.find(qn('w:t'))
                    if t is not None:
                        t.text = ''
        new_trs.append(new_tr)

    parent = template_tr.getparent()
    insert_idx = list(parent).index(template_tr)
    parent.remove(template_tr)
    for offset, new_tr in enumerate(new_trs):
        parent.insert(insert_idx + offset, new_tr)

doc.save(sys.argv[4])